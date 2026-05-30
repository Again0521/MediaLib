from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTLINE_DOC = ROOT / "MediaLIB_系统概要设计.docx"
DETAIL_DOC = ROOT / "MediaLIB_详细设计.docx"
TODAY = "2026-05-26"

PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "5D6B78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F7FB"
CALLOUT_FILL = "EEF6FF"
BORDER = "B9C7D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_text_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_run_font(run, east_asia: str = "PingFang SC") -> None:
    run.font.name = "Calibri"
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header_p = section.header.paragraphs[0]
    header_p.text = title
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    for run in header_p.runs:
        set_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = f"MediaLIB · {TODAY}"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_after = Pt(0)
    for run in footer_p.runs:
        set_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="C6DAF0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}：")
    set_run_font(label_run)
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    body_run = p.add_run(text)
    set_run_font(body_run)
    body_run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_text_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run)


def build_outline_doc() -> None:
    doc = Document()
    configure_doc(doc, "MediaLIB 系统概要设计")
    add_title(
        doc,
        "MediaLIB 系统概要设计",
        "版本：2026-05-26；适用范围：macOS 原生媒体库、播放、扫描、UI、音乐标签编辑与打包架构。",
    )
    add_callout(
        doc,
        "设计目标",
        "MediaLIB 面向本地硬盘、移动硬盘、NAS 挂载目录、SMB/FTP 网络设备和 Emby 服务，提供视频与音乐并列的一体化媒体库。系统优先保持原生体验、可维护小步迭代、离线可用和高性能滚动交互。",
    )

    doc.add_heading("1. 系统定位", level=1)
    add_para(
        doc,
        "MediaLIB 是 SwiftUI + AppKit 混合实现的 macOS 桌面应用。核心边界是管理媒体索引、展示分类与详情、驱动本地/远程播放、同步 Emby 资源、维护设置与隐私保险库。应用不移动用户媒体文件，分类与重分类均写入内部索引。",
    )
    add_bullets(
        doc,
        [
            "用户可见名称统一为 MediaLIB；内部 bundle id、Swift target、数据库目录与 UserDefaults key 继续沿用 MediaLib 以兼容旧数据。",
            "视频播放使用应用内 libmpv render API；音乐播放使用 AVFoundation AVPlayer，同一控制器在展开页和底部迷你播放器之间保持。",
            "音乐标签编辑使用原生 MusicTag 精简工作台，默认只更新 MediaLIB 索引，用户显式打开写入文件后才修改本地音频标签。",
            "普通页面统一使用白色液态玻璃、象牙白/珍珠冷白左上环境光、卡片边缘轮廓光；音乐展开页例外，使用专辑封面取色沉浸背景。",
        ],
    )

    doc.add_heading("2. 上下文与外部依赖", level=1)
    add_table(
        doc,
        ["外部对象", "用途", "交互方式"],
        [
            ["本地/NAS 文件系统", "递归扫描媒体、读取海报、生成缩略图、播放本地文件", "FileManager、AVFoundation、ffmpeg 兜底、libmpv/AVPlayer"],
            ["SMB/FTP 网络设备", "通过 macOS 挂载后纳入扫描", "系统挂载与目录选择，应用复用本地扫描器"],
            ["Emby Server", "同步 Movie/Series/Episode/Audio 资源与远程播放地址", "Emby API；token 存 Keychain；externalID 建立稳定映射"],
            ["TMDB / MusicBrainz / iTunes", "按用户触发补全视频或音乐元数据", "网络搜索服务；用户配置 API 或数据源"],
            ["ffmpeg", "MKV 截帧兜底与显式音乐标签写回", "App 包内或系统路径可执行文件；写回先生成临时文件再替换"],
            ["macOS 系统能力", "Touch ID、AirPlay、媒体键、窗口、Material", "LocalAuthentication、AVKit、MediaPlayer、AppKit/SwiftUI"],
        ],
        [1800, 3780, 3780],
    )

    doc.add_heading("3. 总体分层", level=1)
    add_table(
        doc,
        ["层级", "主要文件/模块", "职责"],
        [
            ["App 层", "MediaLibApp.swift、ContentView.swift、AppState.swift", "应用生命周期、全局状态、主导航、播放器覆盖层、扫描与库刷新编排"],
            ["UI 层", "Views/*、AppColors.swift、MusicPlayerView.swift、PlayerView.swift", "页面、卡片、筛选、详情、播放界面、玻璃材质和交互动效"],
            ["播放层", "MpvPlayerController、LibMpvClient、AirPlayRoutePickerButton", "视频 libmpv 渲染、音乐 AVPlayer、AirPlay 路由、本机同播、系统媒体控制"],
            ["Core 服务层", "MediaScanner、ThumbnailGenerator、AudioMetadataReader、MusicTagEditingService、MetadataSearchService、EmbyService", "扫描解析、封面/缩略图、音频标签读写、远程元数据、Emby 同步"],
            ["数据层", "DatabaseManager、MediaRepository、SourceRepository、Models", "SQLite 表结构访问、媒体条目、媒体源、设置、稳定 ID 与日期编码"],
            ["交付层", "scripts/package_dmg.sh、generate_icon.swift", "生成图标、构建 app、复制依赖、打包 DMG"],
        ],
        [1560, 2860, 4940],
    )

    doc.add_heading("4. 核心数据流", level=1)
    add_numbered(
        doc,
        [
            "用户添加媒体源并选择分类；本地/NAS/已挂载网络目录进入 MediaScanner，Emby 源进入 EmbyService。",
            "扫描阶段解析路径、文件名、音频标签和本地元数据，生成 ParsedMediaFile，再写入 MediaRepository。",
            "AppState.reload() 重建媒体派生缓存、首页统计、父子映射、音乐列表候选和健康提示缓存，并推进 libraryRevision。",
            "页面根据目标分类、搜索、筛选、排序和 libraryRevision 刷新本地快照；普通媒体库使用 LibrarySnapshotCache，音乐库使用预计算 MusicTrackRowModel。",
            "播放入口只设置 activePlayerItem；视频由独立 SwiftUI 播放窗口承载 libmpv，音乐由 ContentView 覆盖层中的 MusicPlaybackHost 和 MusicMiniPlayerBar 承载。",
            "MusicTag 工作台在设置页内按范围批量匹配音乐元数据，用户预览编辑后增量更新索引；显式写入文件时通过 ffmpeg 生成临时文件并安全替换。",
        ],
    )

    doc.add_heading("5. 关键质量属性", level=1)
    add_table(
        doc,
        ["质量属性", "设计策略", "当前约束"],
        [
            ["性能", "派生缓存、快照缓存、异步文件健康检查、hover 节流、滚动时抑制指针光效", "列表 body 不做全量 filter/sort/reduce 或同步 FileManager 检查"],
            ["视觉一致性", "普通页面统一 SurfaceBackground/AppPageBackground；控件复用 GlassCapsuleControl 和玻璃表单", "除音乐展开页外保留低饱和象牙白/珍珠冷白左上光和边缘光"],
            ["播放可靠性", "视频 libmpv render API；音乐复用 AVPlayer；AirPlay 路由由可见 SwiftUI 按钮触发原生选择器", "不要恢复完整 mpv 子窗口或音乐独立 NSWindow"],
            ["隐私", "保险库锁定时隐藏路径、文件名、扫描中文件名和改密入口", "分类只改内部索引，不移动用户文件"],
            ["文件安全", "MusicTag 默认只改索引，显式写入才修改音频文件，远程/不可写/不支持格式逐条失败", "写回必须临时文件成功后替换，不能后台静默批量改原文件"],
            ["可交付性", "SwiftPM 构建、MediaLibChecks、DMG 打包固定命令", "每次交付重新生成 dist/MediaLib.dmg"],
        ],
        [1560, 4540, 3260],
    )

    doc.add_heading("6. 当前版本重点变更", level=1)
    add_bullets(
        doc,
        [
            "音乐底栏 hover 光效扩展到完整 72pt 控制条，高光与专辑色柔光被裁剪在底栏圆角内部，专辑封面不再被左侧圆角裁切。",
            "普通列表滚轮/触控板滑动期间启用 PointerScrollActivityMonitor，临时关闭指针液态光、检视倾斜和 hover 放大，拖动停止后恢复完整效果。",
            "AirPlay 本机同播改为主 AVPlayer 保持本机输出，另建音频路由代理 AVPlayer 交给 AVRoutePickerView 选择外部设备，减少同播设置不生效的概率。",
            "设置页新增 MusicTag 精简工作台，避免引入 GPL/额外限制的 Web 工程，用原生 Swift 弹层和 ffmpeg 写回核心完成批量刮削、编辑和显式标签写入。",
        ],
    )

    doc.add_heading("7. 风险与后续关注", level=1)
    add_bullets(
        doc,
        [
            "AirPlay 行为受系统路由状态回调影响，仍需在真实设备上验证外部播放激活、暂停、seek、切歌和关闭路径。",
            "SwiftUI 玻璃材质、blur、shadow 和 hover 动效容易进入滚动热路径，后续新增列表卡片应默认使用 staticSurfaceBackground。",
            "libmpv、ffmpeg 和 Python/动态库打包依赖需要持续在干净机器上验证加载路径。",
        ],
    )
    doc.save(OUTLINE_DOC)


def build_detail_doc() -> None:
    doc = Document()
    configure_doc(doc, "MediaLIB 详细设计")
    add_title(
        doc,
        "MediaLIB 详细设计",
        "版本：2026-05-26；覆盖模块：数据、扫描、UI、播放、AirPlay、MusicTag、本轮性能修复与验证。",
    )
    add_callout(
        doc,
        "实现原则",
        "所有设计以现有代码小步迭代为前提。普通页面坚持统一白色玻璃语言和低饱和象牙白/珍珠冷白环境光；音乐展开播放器保持专辑色沉浸体验；长列表热路径优先减少高频状态写入和离屏合成。",
    )

    doc.add_heading("1. 模块职责", level=1)
    add_table(
        doc,
        ["模块", "输入", "输出/状态", "关键约束"],
        [
            ["AppState", "数据库、设置、扫描结果、用户操作", "items、派生缓存、musicQueue、activePlayerItem、libraryRevision", "reload 后集中生成缓存；清除播放记录不触发整库 reload"],
            ["LibraryView", "SidebarDestination、搜索、筛选、排序", "普通媒体快照与海报墙", "使用 LibrarySnapshotCache；切换目标显式刷新并标记 destination"],
            ["MusicLibraryView", "MusicLibrarySection、搜索、歌词缓存、排序", "歌曲/专辑/艺术家等快照", "MusicTrackRowModel 预计算；滚轮滑动期间抑制 hover 动效"],
            ["MusicPlayerView", "activePlayerItem、播放器状态、专辑色板", "展开播放页、歌词卡片、控制栏", "不按歌曲 id 强制重建；歌词解析缓存到状态"],
            ["MusicTagScraperSheet", "音乐范围、数据源、用户编辑稿、写入开关", "批量候选列表、展开式标签编辑、索引/文件写入进度", "使用 LazyVStack 和静态玻璃卡片；默认不改原文件"],
            ["MusicTagEditingService", "MusicTagDraft、本地音频文件、ffmpeg", "写入报告或逐条错误", "先写临时文件并校验，再替换原文件；远程/不可写/不支持格式失败"],
            ["MpvPlayerController", "播放项、设置、用户控制", "libmpv/AVPlayer 状态、AirPlay 路由、音量与进度", "视频与音乐路径分离；音乐切歌复用 AVPlayer"],
            ["MediaScanner", "媒体源路径和分类", "规范化媒体条目", "排除歌词、字幕、图片、cue、nfo 等旁路文件；音乐按路径去重"],
        ],
        [1680, 2220, 2820, 2640],
    )

    doc.add_heading("2. UI 设计", level=1)
    add_para(
        doc,
        "AppColors.swift 是普通页面设计系统的中心。AppPageBackground 提供窗口级象牙白/珍珠冷白离屏环境光，LiquidGlassSurfaceLayer 提供白色厚玻璃、静态边缘光和可选 pointer 光源。滚动热路径优先使用 staticSurfaceBackground，只有需要鼠标局部光源的交互卡片使用 surfaceBackground。",
    )
    add_table(
        doc,
        ["组件", "材质/行为", "性能策略"],
        [
            ["AppPageBackground", "普通页面左上象牙白/珍珠冷白环境光、偏白玻璃底", "静态背景，不跟随鼠标高频刷新"],
            ["SurfaceBackground", "玻璃卡片，可传递卡片内鼠标光源给按钮", "仅在需要 pointer 响应的局部使用"],
            ["staticSurfaceBackground", "保留左上染色和边缘轮廓光，无连续 pointer 监听", "长列表、海报、设置分组默认路径"],
            ["GlassCapsuleControl", "首页、视频筛选、音乐筛选共用胶囊", "减少重复实现和样式漂移"],
            ["PosterCardView", "封面区域检视倾斜与 hover 放大", "滚轮/触控板滑动时清空 hoverActive"],
            ["MusicSongRow", "当前 hover 行轻微放大、高光、描边", "滚动时临时关闭 hover 反馈，停止后恢复"],
        ],
        [2020, 4080, 3260],
    )

    doc.add_heading("3. 本轮底栏光效修复", level=1)
    add_para(
        doc,
        "MusicMiniPlayerBar 的视觉问题来自内容垂直 padding 位于玻璃裁剪层外侧，导致鼠标光效和专辑色柔光只覆盖内层内容高度。修复后将横向与纵向 padding 移入同一个圆角 ZStack/HStack，外层保持 72pt 固定高度，MusicMiniAlbumGlowLayer、玻璃表面和 pointerLiquidLight 使用同一圆角裁剪范围。",
    )
    add_bullets(
        doc,
        [
            "底栏仍由 ContentView 的窗口级 GeometryReader 贴底，宽度按侧栏状态弹性计算。",
            "封面、曲目信息、播放按钮、进度和工具按钮都在完整玻璃底栏内部布局，避免左侧专辑封面被圆角切掉。",
            "专辑色柔光只裁剪在底栏内部，主要用于边缘光和按钮轮廓，不向列表区域溢出。",
        ],
    )

    doc.add_heading("4. 滚轮与触控板性能修复", level=1)
    add_para(
        doc,
        "拖动系统滚动条较顺滑而滚轮/触控板卡顿，主要因为指针停留在列表 cell 上时，wheel/trackpad 事件会持续触发行 hover 状态、检视倾斜、液态光 GeometryReader 和动画合成。修复策略是不牺牲最终视觉效果，只在实际滚动事件发生后的短时间窗口内抑制 hover 动效。",
    )
    add_numbered(
        doc,
        [
            "在 AppColors.swift 新增 PointerScrollActivityMonitor，使用 AppKit 本地 scrollWheel 事件监控并判断事件是否发生在当前 ScrollView bounds 内。",
            "新增 suppressHoverEffectsDuringScroll() 修饰器，把滚动活动写入环境值 suppressPointerHoverDuringScroll，180ms 无滚动后恢复。",
            "pointerLiquidLight、pointerLiquidEdge 和 pointerInspectTilt 读取环境值，滚动期间直接降级为原内容或清空指针位置。",
            "PosterCardView、MusicSongRow、SourceRowView 和 StatTile 使用 hoverActive = isHovering && !suppressPointerHoverDuringScroll，并在滚动开始时清理 hover 状态。",
            "在 LibraryView、MusicLibraryView、HomeView、SourcesView、DetailView、SettingsView 的 ScrollView 上接入修饰器，覆盖用户实际滑动路径。",
        ],
    )

    doc.add_heading("5. 音乐 AirPlay 本机同播设计", level=1)
    add_para(
        doc,
        "旧实现依赖同一个 AVPlayer 的 externalPlayback 状态再创建本机镜像，但当系统把输出路线整体切到 AirPlay 时，本机保留输出可能失效。本轮改为双播放器路由：主音乐 AVPlayer 禁止外部播放并保留本机声音；音频路由代理 AVPlayer 绑定同一音频文件、允许外部播放，并交给原生 AVRoutePickerView 选择 AirPlay 设备。",
    )
    add_table(
        doc,
        ["场景", "主 AVPlayer", "音频路由代理 AVPlayer", "预期结果"],
        [
            ["本机同播开启", "allowsExternalPlayback = false，保持本机输出", "allowsExternalPlayback = true，选择外部设备后同步播放", "本机与 AirPlay 设备同时播放"],
            ["本机同播关闭", "允许外部播放，由系统路由接管", "释放/暂停", "遵循系统当前外部输出路线"],
            ["seek/倍速/音量/暂停", "执行用户操作并保存设置", "同步时间、速率、播放状态和音量", "两端状态尽量一致"],
            ["切歌/关闭/失败", "replaceCurrentItem 或 teardown", "清理代理、取消观察、递增 routePickerRevision", "避免上一首代理继续出声"],
        ],
        [1740, 2720, 2940, 1960],
    )

    doc.add_heading("6. 数据与缓存设计", level=1)
    add_table(
        doc,
        ["缓存/状态", "生命周期", "用途"],
        [
            ["libraryRevision", "媒体库数据或派生缓存变化时递增", "驱动页面快照刷新和列表状态同步"],
            ["LibrarySnapshotCache", "按 destination/search/filter/sort/revision 命中", "避免普通媒体库 body 反复全量过滤排序"],
            ["MusicTrackRowModel", "音乐快照刷新时生成", "预计算文件名、时长、艺人、专辑和歌词状态"],
            ["ArtworkImageCache", "进程内 NSCache + 缺失路径短缓存", "减少海报解码和同步读盘；body 只读已缓存比例"],
            ["Lyrics parse cache", "歌词文本变化时刷新", "避免播放进度刷新时重复解析 LRC"],
            ["AirPlay route proxy state", "音乐播放会话内", "保存代理播放器、观察者和 routePickerRevision"],
            ["MusicTag candidates", "MusicTag 弹层会话内", "保存网络匹配结果、编辑草稿、选中状态和写入状态"],
        ],
        [2260, 3200, 3900],
    )

    doc.add_heading("7. 扫描与元数据", level=1)
    add_bullets(
        doc,
        [
            "本地扫描递归穿透嵌套目录，自动识别媒体源同时导入视频和音频；音频文件不受视频 50MB 阈值限制。",
            "音乐标签读取覆盖 common metadata 以及 iTunes/ID3/QuickTime 等格式级 metadata，优先使用内嵌封面和歌词。",
            "MusicTag 工作台复用 MusicBrainz / iTunes Search 做批量候选，先预览并允许用户编辑，再更新 MediaLIB 索引或显式写回音频文件标签。",
            "MusicTag 写回没有引入 GPL/额外限制的 music-tag-web 整包；当前用原生 Swift + ffmpeg 保留核心标签编辑能力，降低许可证、包体和运行时内存风险。",
            "扫描器排除 `.lrc`、`.txt`、字幕、本地图片、`.cue`、`.nfo` 等旁路文件，只导入规范化后的真实媒体文件。",
            "Emby 登录同步通过 sourcePath + externalID 维持与 Emby ItemId 的映射，token 存系统钥匙串。",
        ],
    )

    doc.add_heading("8. 错误处理与安全", level=1)
    add_table(
        doc,
        ["风险", "处理策略"],
        [
            ["NAS 或文件提供器路径不可达", "列表路径不做同步存在性检查；后台 file health 任务更新状态，详情/预览/播放再按需检查"],
            ["libmpv 或 OpenGL 初始化异常", "OpenGL pixel format 创建失败时使用标准 NSOpenGLView 兜底；播放器失败时走 errorMessage"],
            ["AirPlay 状态回调滞后", "路线选择后延长探测并同步代理播放状态；切歌和展开/收起时刷新路线"],
            ["MusicTag 写入失败", "逐条显示失败原因；临时文件未成功时不替换原文件；封面嵌入失败可退回只写文字标签"],
            ["保险库锁定泄露路径", "锁定状态过滤保险库条目、扫描文件名、媒体源路径和改密入口"],
            ["元数据网络失败", "搜索和补全为用户触发动作，失败不破坏本地媒体索引"],
        ],
        [2760, 6600],
    )

    doc.add_heading("9. 验证方案", level=1)
    add_bullets(
        doc,
        [
            "构建：env DEVELOPER_DIR=/Applications/Xcode.app/Contents/Developer swift build。",
            "健康检查：env DEVELOPER_DIR=/Applications/Xcode.app/Contents/Developer swift run MediaLibChecks。",
            "交付：scripts/package_dmg.sh 生成 dist/MediaLib.dmg。",
            "MusicTag 回归：打开设置页 MusicTag 弹层，验证范围切换、匹配进度、展开编辑、默认仅更新索引和写入开关提示；真实文件写回需用测试音频验证。",
            "人工回归：音乐列表用鼠标滚轮和触控板快速滑动，停止后 hover 放大和光效恢复；底栏 hover 光效覆盖完整高度；AirPlay 本机同播在真实外部设备上验证。",
            "文档 QA：使用 Documents 技能渲染两个 DOCX 为 PNG，检查标题、表格、分页和中文字体无裁切或重叠。",
        ],
    )

    doc.add_heading("10. 后续维护规则", level=1)
    add_bullets(
        doc,
        [
            "新增普通页面卡片时优先复用 staticSurfaceBackground，只有明确需要局部鼠标光源时使用 surfaceBackground。",
            "不要把音乐展开页材质或专辑色光效扩散到普通页面；普通页面保留偏白、低噪声的象牙白/珍珠冷白玻璃。",
            "音乐播放、切歌、队列和 AirPlay 继续以单一 MpvPlayerController 为状态源；不要恢复独立音乐窗口。",
            "后续扩展 MusicTag 时继续保持默认不写文件、显式写入、逐条失败和许可证兼容边界；不要把 GPL Web 工程直接塞进 App。",
            "任何 UI/功能改动都同步更新 README、CHANGELOG、handoff、开发说明并重新生成 DMG。",
        ],
    )
    doc.save(DETAIL_DOC)


def main() -> None:
    build_outline_doc()
    build_detail_doc()
    print(OUTLINE_DOC)
    print(DETAIL_DOC)


if __name__ == "__main__":
    main()
